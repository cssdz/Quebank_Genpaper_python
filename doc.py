import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor
import mysql_op
from docx.oxml.ns import qn

dbManager = mysql_op.DBManager()


class CreatePaper:
    def __init__(self, semester, dept, course, class_, s_name, s_no, report_name, que, path):
        self.semester = semester
        self.dept = dept
        self.course = course
        self.class_ = class_
        self.s_name = s_name
        self.s_no = s_no
        self.report_name = report_name
        self.doc_ = None
        self.que = que
        self.path = path

    def create_format(self):
        # .docx正文格式
        self.doc_ = docx.Document()
        self.doc_.styles['Normal'].font.name = u'宋体'
        self.doc_.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.doc_.styles['Normal'].font.size = Pt(12.5)
        self.doc_.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

        # 创建试卷标题及格式
        title = "%s南阳理工学院%s考试" % (self.semester, self.dept)
        head = self.doc_.add_heading(title, level=0)
        head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 创建学生信息
        stu_info = ('课程:{0}' + '  ' + '班级:{1}' + '  ' + '姓名:{2}' + '  ' + '学号:{3}').format(self.course, self.class_,
                                                                                           self.s_name, self.s_no)
        self.doc_.add_paragraph(stu_info)
        self.doc_.save(r"%s" % self.path + "%s" % self.report_name)

    def save_que(self):
        type_no = ("一", "二", "三", "四", "五", "六", "七", "八")
        sql = "SELECT * FROM no_type;"
        info = dbManager.fetchall(sql)
        que_, que_title = 0, ""
        for i in range(len(info)):
            type_info = type_no[i] + "." + str(info[i][1])
            self.doc_.add_heading(type_info, level=1)
            for j in range(len(self.que[i])):
                factors, option = len(self.que[i][j]), ""
                for k in range(factors):
                    if k == factors - 1:
                        continue
                    if k == 0:
                        que_ += 1
                        que_title = str(que_) + "." + self.que[i][j][k]
                    else:
                        option += str(self.que[i][j][k]) + "  \n"

                self.doc_.add_paragraph(que_title)
                self.doc_.add_paragraph(option)
        self.doc_.save(r"%s" % self.path + "%s" % self.report_name)

    def combine_file(self):
        self.create_format()
        self.save_que()
